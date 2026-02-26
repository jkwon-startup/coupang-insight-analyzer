"""Word 리포트 내보내기"""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordExporter:
    def generate(
        self,
        product_data,
        story_result,
        review_result,
        qna_result,
        full_result,
    ) -> bytes:
        doc = Document()

        # 제목
        title = doc.add_heading("Coupang Insight Analyzer", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("쿠팡 상품 분석 리포트")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()

        # 상품 기본 정보
        if product_data:
            doc.add_heading("상품 기본 정보", level=1)
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            info = [
                ("상품명", product_data.get("title", "")),
                ("가격", product_data.get("price", "")),
                ("리뷰 수", str(product_data.get("review_count", ""))),
                ("URL", product_data.get("url", "")),
            ]
            for label, value in info:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
                row.cells[1].text = value

            doc.add_paragraph()

        # 스토리 분석
        if story_result:
            doc.add_heading("상세페이지 스토리 분석", level=1)
            self._add_markdown_content(doc, story_result)

        # 리뷰 분석
        if review_result:
            doc.add_heading("리뷰 분석", level=1)
            self._add_markdown_content(doc, review_result)

        # Q&A 분석
        if qna_result:
            doc.add_heading("상품문의(Q&A) 분석", level=1)
            self._add_markdown_content(doc, qna_result)

        # 종합 리포트
        if full_result:
            doc.add_heading("종합 리포트", level=1)
            self._add_markdown_content(doc, full_result)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _add_markdown_content(self, doc, text: str):
        """마크다운 텍스트를 Word 포맷으로 변환하여 추가"""
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
            elif line.startswith("| "):
                # 테이블은 일반 텍스트로 처리
                doc.add_paragraph(line, style="Normal")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)
